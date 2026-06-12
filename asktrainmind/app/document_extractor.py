"""
Document content extraction for Phase 3.

Extracts per-page text and embedded images from downloaded documents.
Primary format: PDF via PyMuPDF (fitz) — imported lazily so the app works
without it (offline-safe). Office formats (docx/pptx/xlsx) are best-effort
via lightweight libraries if available; deferred when not.

Extension hook for OCR is documented below (deferred to Phase 4).
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path


# ---------------------------------------------------------------------------
# Public dataclasses
# ---------------------------------------------------------------------------

@dataclass
class PageText:
    page_number: int  # 1-based
    text: str


@dataclass
class DocImage:
    page_number: int
    mime_type: str
    data: bytes


@dataclass
class ExtractedDocument:
    source_url: str
    local_path: Path
    pages: list[PageText] = field(default_factory=list)
    images: list[DocImage] = field(default_factory=list)
    page_count: int = 0
    status: str = "ok"
    message: str = ""


# ---------------------------------------------------------------------------
# PDF extraction via PyMuPDF
# ---------------------------------------------------------------------------

def _extract_pdf(source_url: str, local_path: Path) -> ExtractedDocument:
    """Extract text and images from a PDF file using PyMuPDF (fitz)."""
    try:
        import fitz  # type: ignore[import-untyped]
    except ImportError:
        return ExtractedDocument(
            source_url=source_url,
            local_path=local_path,
            status="missing_dependency",
            message=(
                "PyMuPDF (pymupdf) non è installato. "
                "Installare con: pip install pymupdf. "
                "Estrazione testo PDF non disponibile."
            ),
        )

    try:
        pdf = fitz.open(str(local_path))
    except Exception as exc:
        return ExtractedDocument(
            source_url=source_url,
            local_path=local_path,
            status="open_error",
            message=f"Impossibile aprire il PDF: {exc}",
        )

    pages: list[PageText] = []
    images: list[DocImage] = []

    try:
        for page_index in range(pdf.page_count):
            page = pdf[page_index]
            page_num = page_index + 1

            # Text
            text = page.get_text()
            pages.append(PageText(page_number=page_num, text=text))

            # Embedded images
            image_list = page.get_images(full=True)
            for img_info in image_list:
                xref = img_info[0]
                try:
                    base_image = pdf.extract_image(xref)
                    img_bytes = base_image.get("image", b"")
                    img_ext = base_image.get("ext", "png").lower()
                    mime = f"image/{img_ext}" if img_ext != "jpeg" else "image/jpeg"
                    if img_bytes:
                        images.append(DocImage(page_number=page_num, mime_type=mime, data=img_bytes))
                except Exception:
                    pass  # Skip corrupt embedded images

        return ExtractedDocument(
            source_url=source_url,
            local_path=local_path,
            pages=pages,
            images=images,
            page_count=pdf.page_count,
        )
    except Exception as exc:
        return ExtractedDocument(
            source_url=source_url,
            local_path=local_path,
            pages=pages,
            images=images,
            page_count=getattr(pdf, "page_count", len(pages)),
            status="partial_error",
            message=f"Estrazione parziale: {exc}",
        )
    finally:
        pdf.close()


# ---------------------------------------------------------------------------
# Office format extraction (best-effort, optional)
# ---------------------------------------------------------------------------

def _extract_docx(source_url: str, local_path: Path) -> ExtractedDocument:
    """Extract text from a DOCX file via python-docx if available."""
    try:
        from docx import Document as DocxDocument  # type: ignore[import-untyped]
    except ImportError:
        return ExtractedDocument(
            source_url=source_url,
            local_path=local_path,
            status="missing_dependency",
            message=(
                "python-docx non è installato. "
                "Estrazione testo DOCX non disponibile (formato opzionale)."
            ),
        )

    try:
        doc = DocxDocument(str(local_path))
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        pages = [PageText(page_number=1, text=full_text)]
        return ExtractedDocument(
            source_url=source_url,
            local_path=local_path,
            pages=pages,
            page_count=1,
        )
    except Exception as exc:
        return ExtractedDocument(
            source_url=source_url,
            local_path=local_path,
            status="open_error",
            message=f"Impossibile aprire il DOCX: {exc}",
        )


def _extract_unsupported(source_url: str, local_path: Path) -> ExtractedDocument:
    return ExtractedDocument(
        source_url=source_url,
        local_path=local_path,
        status="unsupported_format",
        message=(
            f"Formato non supportato: {local_path.suffix or 'sconosciuto'}. "
            "Solo PDF è supportato in modo completo. "
            "DOCX è opzionale (richiede python-docx). "
            "PPTX/XLSX non sono supportati in questa fase."
        ),
    )


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

def extract_document(source_url: str, local_path: Path) -> ExtractedDocument:
    """
    Extract text and images from a downloaded document.

    Dispatches to the appropriate extractor based on the file extension.
    Always returns an ExtractedDocument — never raises.
    """
    suffix = local_path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(source_url, local_path)
    if suffix == ".docx":
        return _extract_docx(source_url, local_path)
    # Other Office formats: pptx, xlsx, doc, ppt, xls — not supported yet
    return _extract_unsupported(source_url, local_path)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def get_page_text(doc: ExtractedDocument, page_number: int) -> str:
    """Return the extracted text for the given 1-based page number, or ''."""
    for page in doc.pages:
        if page.page_number == page_number:
            return page.text
    return ""


def find_text_snippet(doc: ExtractedDocument, query: str, max_chars: int = 500) -> str:
    """
    Find the best snippet in the document that contains the query string.

    Searches all pages for the query (case-insensitive).
    Returns up to *max_chars* characters of context around the first match,
    or the first *max_chars* chars of page 1 if no match found.
    """
    if not query:
        if doc.pages:
            return doc.pages[0].text[:max_chars]
        return ""

    query_lower = query.lower()
    for page in doc.pages:
        text = page.text
        idx = text.lower().find(query_lower)
        if idx >= 0:
            start = max(0, idx - 100)
            end = min(len(text), idx + len(query) + 300)
            snippet = text[start:end]
            if len(snippet) > max_chars:
                snippet = snippet[:max_chars]
            return snippet

    # No match — return beginning of first page
    if doc.pages:
        return doc.pages[0].text[:max_chars]
    return ""


# ---------------------------------------------------------------------------
# OCR hook (deferred to Phase 4)
# ---------------------------------------------------------------------------
# To add OCR for scanned PDFs, implement _extract_pdf_ocr() here using
# pytesseract or easyocr. Call it from extract_document() when PyMuPDF
# returns empty text for all pages. Deferred to Phase 4.
